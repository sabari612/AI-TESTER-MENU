from io import BytesIO

from processing.normalize_menu import parse_menu_text


def read_docx_menu(file_bytes):
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required for Word extraction.") from exc

    document = Document(BytesIO(file_bytes))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" - ".join(cells))
    return parse_menu_text("\n".join(lines))

